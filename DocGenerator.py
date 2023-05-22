import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import pandas as pd
from docx import Document
import os


def select_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xls")])
    excel_file_entry.delete(0, tk.END)
    excel_file_entry.insert(0, file_path)


def select_word_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    word_file_entry.delete(0, tk.END)
    word_file_entry.insert(0, file_path)


def select_output_folder():
    folder_path = filedialog.askdirectory()
    output_folder_entry.delete(0, tk.END)
    output_folder_entry.insert(0, folder_path)


def generate_documents():
    excel_file = excel_file_entry.get()
    word_file = word_file_entry.get()
    suffix_column = suffix_column_entry.get()
    output_folder = output_folder_entry.get()

    if not excel_file or not word_file or not output_folder:
        messagebox.showerror("Error", "Please select Excel file, Word file, and output folder.")
        return

    try:
        data = pd.read_excel(excel_file)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to read Excel file:\n{e}")
        return

    if suffix_column not in data.columns:
        messagebox.showerror("Error", "Invalid suffix column selected.")
        return

    combined_document = None
    for index, row in data.iterrows():
        document = Document(word_file)
        for column in data.columns:
            placeholder = '{' + column + '}'
            value = str(row[column])
            for paragraph in document.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        suffix = str(row[suffix_column])
        filename = f"{index}_{suffix}.docx"
        filepath = os.path.join(output_folder, filename)
        document.save(filepath)

        if combine_var.get():
            if not combined_document:
                combined_document = document
            else:
                for element in document.element.body:
                    combined_document.element.body.append(element)

    if combine_var.get() and combined_document:
        combined_filename = os.path.join(output_folder, "combined.docx")
        combined_document.save(combined_filename)

    messagebox.showinfo("Success", "Documents generated successfully.")


# Create the main window
window = tk.Tk()
window.title("Document Generator")

# Excel file selection
excel_file_label = tk.Label(window, text="Select Excel File:")
excel_file_label.pack()
excel_file_entry = tk.Entry(window, width=50)
excel_file_entry.pack()
excel_file_button = tk.Button(window, text="Browse", command=select_excel_file)
excel_file_button.pack()

# Word file selection
word_file_label = tk.Label(window, text="Select Word File:")
word_file_label.pack()
word_file_entry = tk.Entry(window, width=50)
word_file_entry.pack()
word_file_button = tk.Button(window, text="Browse", command=select_word_file)
word_file_button.pack()

# Suffix column selection
suffix_column_label = tk.Label(window, text="Select Suffix Column:")
suffix_column_label.pack()
suffix_column_entry = tk.Entry(window)
suffix_column_entry.pack()

# Output folder selection
output_folder_label = tk.Label(window, text="Select Output Folder:")
output_folder_label.pack()
output_folder_entry = tk.Entry(window, width=50)
output_folder_entry.pack()
output_folder_button = tk.Button(window, text="Browse", command=select_output_folder)
output_folder_button.pack()

# Combine documents option
combine_var = tk.BooleanVar()
combine_var.set(False)
combine_checkbutton = tk.Checkbutton(window, text="Combine generated documents into a single file", variable=combine_var)
combine_checkbutton.pack()

# Generate button
generate_button = tk.Button(window, text="Generate Documents", command=generate_documents)
generate_button.pack()

# Run the main event loop
window.mainloop()
def generate_documents():
    excel_file = excel_file_entry.get()
    word_file = word_file_entry.get()
    suffix_column = suffix_column_entry.get()
    output_folder = output_folder_entry.get()

    if not excel_file or not word_file or not output_folder:
        messagebox.showerror("Error", "Please select Excel file, Word file, and output folder.")
        return

    try:
        data = pd.read_excel(excel_file)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to read Excel file:\n{e}")
        return

    if suffix_column not in data.columns:
        messagebox.showerror("Error", "Invalid suffix column selected.")
        return

    # Create an empty Document object to hold the combined document
    combined_document = Document()

    # Loop through each row in the Excel sheet
    for index, row in data.iterrows():

        # Load the base Word document
        document = Document(word_file)

        # Replace placeholders with values from the current row in the Excel sheet
        for column in data.columns:
            placeholder = '{' + column + '}'
            value = str(row[column])
            for paragraph in document.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

        # Generate a filename for the output document based on the current row index and the suffix column
        suffix = str(row[suffix_column])
        filename = f"{index}_{suffix}.docx"
        filepath = os.path.join(output_folder, filename)

        # Save the current document to disk
        document.save(filepath)

        # If combine checkbox is checked, add contents of each document to combined document with page break
        if combine_var.get():
            # Load the current document from disk and append its contents to the combined document
            with open(filepath, 'rb') as f:
                sub_document = Document(f)
                if index > 0:
                    combined_document.add_page_break()  # Add a page break before appending the new sub-document
                for element in sub_document.element.body:
                    combined_document.element.body.append(element)
     

    # If the 'combine_var' checkbox is checked, save the combined document to disk
    if combine_var.get() and len(data) > 0:
        combined_filename = os.path.join(output_folder, "combined.docx")
        combined_document.save(combined_filename)

    messagebox.showinfo("Success", "Documents generated successfully.")
